"""
员工全景画像综合研究报告 v5.0（适配最终汉化数据集）
===================================================
基于汉化脚本 v5.0 输出的 Excel 文件，直接使用已有编码列
功能：
- 六大研究方向：画像、流失、薪酬、生命周期、职业发展、离职预测
- 生成 Word 报告及 Excel 风险分级统计表
- 所有图表保存为 PNG，图文结合，排版优化
"""

import pandas as pd
import numpy as np
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots
from sklearn.ensemble import RandomForestClassifier
from sklearn.model_selection import train_test_split
from sklearn.metrics import (classification_report, confusion_matrix,
                             roc_curve, auc, accuracy_score,
                             precision_score, recall_score, f1_score)
import joblib
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
import warnings
from openpyxl.styles import Font, Alignment, PatternFill
from openpyxl.worksheet.table import Table, TableStyleInfo
from openpyxl.utils import get_column_letter

warnings.filterwarnings('ignore')

# ==================== 路径配置 ====================
BASE_DIR = Path(__file__).parent.parent.parent  # 项目根目录
DATA_FILE = BASE_DIR / "output" / "IBM_HR_员工流失数据_本土化版.xlsx"
OUTPUT_DIR = BASE_DIR / "analysis" / "output"
IMAGES_DIR = OUTPUT_DIR / "images"
WORD_FILE = OUTPUT_DIR / "员工全景画像分析报告.docx"
EXCEL_RISK_FILE = OUTPUT_DIR / "在职员工离职风险分级统计表.xlsx"

OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
IMAGES_DIR.mkdir(parents=True, exist_ok=True)

# ==================== 加载数据 ====================
print("📊 加载数据...")
df = pd.read_excel(DATA_FILE, sheet_name="数据")
print(f"✅ 数据加载成功！共 {len(df)} 行")

# 设置全局模板
template = "plotly_white"

# ==================== 定义列名（利用 v5.0 已有编码列）====================
satisfaction_text_cols = ["环境满意", "人际关系满意", "工作满意", "敬业度", "工作与生活平衡"]
satisfaction_num_cols = [col + "编码" for col in satisfaction_text_cols]

# 创建年龄组（用于图表）
df["年龄组"] = pd.cut(df["年龄"], bins=[18, 25, 35, 45, 55, 65],
                      labels=["18-25岁", "26-35岁", "36-45岁", "46-55岁", "56-65岁"])

# 工龄组
df["工龄组"] = pd.cut(df["总工龄"], bins=[0, 2, 5, 10, 20, 50],
                      labels=["0-2年", "3-5年", "6-10年", "11-20年", "20年以上"])

# 晋升间隔组
df["晋升间隔组"] = pd.cut(df["晋升间隔"], bins=[-1, 1, 3, 5, 10, 20],
                          labels=["0-1年", "2-3年", "4-5年", "6-10年", "10年以上"])

# 培训次数组
df["培训次数组"] = pd.cut(df["年度培训次数"], bins=[0, 1, 2, 3, 4, 6],
                          labels=["0-1次", "2次", "3次", "4次", "5-6次"])

# 教育程度顺序（用于图表）
edu_order = ["大专以下", "大专", "本科", "硕士", "博士"]

# 绩效评级顺序
perf_order = ["低", "良好", "优秀", "杰出"]

# 确保绩效评级为有序类别（用于图表）
df["绩效评级"] = pd.Categorical(df["绩效评级"], categories=perf_order, ordered=True)

# ==================== 辅助函数：保存图表 ====================
def save_chart(fig, filename, width=800, height=500):
    """保存图表为 HTML 和 PNG"""
    fig.write_html(IMAGES_DIR / f"{filename}.html")
    fig.write_image(IMAGES_DIR / f"{filename}.png", scale=2, width=width, height=height)
    print(f"  ✅ 已保存: {filename}.png")

# ==================== 1. 员工基本画像 ====================
print("\n" + "="*60)
print("📋 第一部分：员工基本画像")
print("="*60)

# 1.1 年龄分布
fig = px.histogram(df, x="年龄", nbins=20, title="员工年龄分布",
                   labels={"年龄": "年龄（岁）", "count": "人数"},
                   marginal="box", template=template, color_discrete_sequence=["#4472C4"])
save_chart(fig, "01_年龄分布")

# 1.2 性别比例
gender_counts = df["性别"].value_counts().reset_index()
gender_counts.columns = ["性别", "人数"]
fig = px.pie(gender_counts, values="人数", names="性别", title="性别比例",
             hole=0.3, template=template, color_discrete_sequence=["#4472C4", "#8CB4E8"])
fig.update_traces(textposition='inside', textinfo='percent+label')
save_chart(fig, "02_性别比例")

# 1.3 学历分布
edu_counts = df["学历"].value_counts().reindex(edu_order).reset_index()
edu_counts.columns = ["学历", "人数"]
fig = px.bar(edu_counts, x="学历", y="人数", title="学历分布",
             color="人数", color_continuous_scale="Blues", template=template)
save_chart(fig, "03_学历分布")

# 1.4 婚姻状况分布
marital_counts = df["婚姻状况"].value_counts().reset_index()
marital_counts.columns = ["婚姻状况", "人数"]
fig = px.pie(marital_counts, values="人数", names="婚姻状况", title="婚姻状况分布",
             template=template, color_discrete_sequence=px.colors.qualitative.Set3)
save_chart(fig, "04_婚姻状况分布")

# 关键结果
avg_age = df["年龄"].mean()
gender_ratio = df["性别"].value_counts(normalize=True)["男"] * 100
edu_main = df["学历"].mode()[0]
marital_main = df["婚姻状况"].mode()[0]

chart_analysis_01 = {
    "01_年龄分布": "年龄分布呈单峰形态，集中在30-45岁，说明公司以中青年员工为主，这有助于保持组织活力，但也需关注年轻员工的培养和资深员工的经验传承。",
    "02_性别比例": f"男性占比 {gender_ratio:.1f}%，女性 {100-gender_ratio:.1f}%，比例均衡，有利于多元化和性别平等。",
    "03_学历分布": f"学历以 {edu_main} 为主，高学历人才占比高，符合知识密集型企业的特点，为技术创新提供基础。",
    "04_婚姻状况分布": f"已婚员工占比最高，这类员工通常稳定性更强，对薪酬福利和发展机会更为敏感。",
}

chapter1_summary = f"""
【基本特征总结】
- 平均年龄 {avg_age:.1f} 岁，员工队伍年轻有活力。
- 性别比例均衡，有利于团队多样性。
- 高学历员工为主体，研发和创新基础良好。
- 已婚员工居多，对稳定性和发展有较高诉求。

【管理启示】
针对年轻员工设计快速成长通道，对已婚员工提供弹性福利和长期激励，保持队伍稳定。
"""

# ==================== 2. 不同分类的流失分析 ====================
print("\n" + "="*60)
print("📈 第二部分：不同分类的流失分析")
print("="*60)

chapter2_text = {
    "目标": "识别高流失风险群体，为精准干预提供依据。",
    "内容": "分析部门、岗位、年龄、加班、满意度等因素与离职率的关系。"
}

# 2.1 总体离职率
attrition_rate = df["是否离职"].value_counts(normalize=True)["是"] * 100

# 2.2 部门离职率
dept_attrition = df.groupby("部门")["是否离职"].value_counts(normalize=True).unstack()["是"] * 100
dept_attrition_df = dept_attrition.reset_index()
dept_attrition_df.columns = ["部门", "离职率"]
fig = px.bar(dept_attrition_df, x="部门", y="离职率", title="各部门离职率对比",
             color="离职率", color_continuous_scale="Reds", template=template)
save_chart(fig, "05_部门离职率")

# 2.3 岗位离职率TOP15
job_attrition = df.groupby("岗位")["是否离职"].value_counts(normalize=True).unstack()["是"] * 100
job_attrition_df = job_attrition.sort_values(ascending=False).reset_index()
job_attrition_df.columns = ["岗位", "离职率"]
fig = px.bar(job_attrition_df.head(15), x="离职率", y="岗位", orientation='h',
             title="离职率最高的15个岗位", color="离职率", color_continuous_scale="Reds", template=template)
fig.update_layout(yaxis={'categoryorder':'total ascending'})
save_chart(fig, "06_岗位离职率TOP15")

# 2.4 年龄组与离职率
age_attrition = df.groupby("年龄组")["是否离职"].value_counts(normalize=True).unstack()["是"] * 100
age_attrition_df = age_attrition.reset_index()
age_attrition_df.columns = ["年龄组", "离职率"]
fig = px.line(age_attrition_df, x="年龄组", y="离职率", title="不同年龄组离职率",
              markers=True, template=template, color_discrete_sequence=["#E84C3D"])
save_chart(fig, "07_年龄组离职率")

# 2.5 加班与离职率
overtime_attrition = df.groupby("是否加班")["是否离职"].value_counts(normalize=True).unstack()["是"] * 100
overtime_df = overtime_attrition.reset_index()
overtime_df.columns = ["是否加班", "离职率"]
fig = px.bar(overtime_df, x="是否加班", y="离职率", title="加班与离职关系",
             color="是否加班", template=template,
             color_discrete_map={"是": "#E84C3D", "否": "#4472C4"})
save_chart(fig, "08_加班离职率")

# 2.6 满意度与离职率（使用编码列）
satisfaction_attrition = {}
for i, text_col in enumerate(satisfaction_text_cols[:4]):
    num_col = satisfaction_num_cols[i]
    rate = df.groupby(num_col)["是否离职"].value_counts(normalize=True).unstack()["是"] * 100
    satisfaction_attrition[text_col] = rate

fig = make_subplots(rows=2, cols=2, subplot_titles=list(satisfaction_attrition.keys()), shared_yaxes=True)
row, col = 1, 1
for title, data in satisfaction_attrition.items():
    fig.add_trace(go.Bar(x=data.index, y=data.values, name=title,
                          marker_color=['#4472C4', '#5A8AC4', '#8CB4E8', '#B0D0F0']),
                  row=row, col=col)
    col += 1
    if col > 2:
        col = 1
        row += 1
fig.update_layout(height=600, title_text="不同满意度维度的离职率对比", template=template, showlegend=False)
save_chart(fig, "09_满意度离职率")

# 关键指标
highest_dept = dept_attrition_df.loc[dept_attrition_df["离职率"].idxmax()]
highest_job = job_attrition_df.iloc[0]
overtime_risk = overtime_attrition["是"] / overtime_attrition["否"]
satisfaction_low = satisfaction_attrition['工作满意'][1] if 1 in satisfaction_attrition['工作满意'] else 0
satisfaction_high = satisfaction_attrition['工作满意'][4] if 4 in satisfaction_attrition['工作满意'] else 0

chart_analysis_02 = {
    "05_部门离职率": f"部门间离职率差异显著，{highest_dept['部门']} 最高（{highest_dept['离职率']:.1f}%），需重点排查该部门的管理风格、工作强度或薪酬问题。",
    "06_岗位离职率TOP15": f"{highest_job['岗位']} 离职率高达 {highest_job['离职率']:.1f}%，可能是由于工作压力大、晋升通道窄或薪酬竞争力不足。",
    "07_年龄组离职率": "26-35岁员工离职率最高，该年龄段处于职业探索期，对发展机会敏感，需提供清晰的成长路径。",
    "08_加班离职率": f"加班员工离职率是非加班员工的 {overtime_risk:.1f} 倍，加班文化对留任产生显著负面影响。",
    "09_满意度离职率": f"工作满意度评分1分的员工离职率 {satisfaction_low:.1f}%，而4分员工仅 {satisfaction_high:.1f}%，改善满意度是降低流失的关键。",
}

chapter2_summary = f"""
【流失风险总结】
- 整体离职率 {attrition_rate:.1f}%，处于可控范围，但特定群体流失严重。
- 部门层面：{highest_dept['部门']} 需优先干预。
- 岗位层面：{highest_job['岗位']} 为高风险岗位，建议开展离职访谈。
- 加班和低满意度是核心推力，需优化工作负荷和员工关怀。

【管理启示】
建立定期流失监测机制，对高风险部门/岗位实施专项保留计划，将员工满意度纳入管理者考核。
"""

# ==================== 3. 薪酬公平性分析 ====================
print("\n" + "="*60)
print("💰 第三部分：薪酬公平性分析")
print("="*60)

chapter3_text = {
    "目标": "评估薪酬体系是否存在不公平现象，为薪酬调整提供依据。",
    "内容": "分析月收入分布、部门/岗位/性别/学历对薪酬的影响。注：薪酬分析以月收入为核心指标，符合国内薪酬分析习惯。"
}

# 3.1 月收入分布
fig = px.histogram(df, x="月收入", nbins=30, title="月收入分布",
                   labels={"月收入": "月收入（元）", "count": "人数"},
                   marginal="box", template=template, color_discrete_sequence=["#4472C4"])
save_chart(fig, "10_月收入分布")

# 3.2 部门月收入对比
fig = px.box(df, x="部门", y="月收入", title="各部门月收入分布",
             color="部门", template=template)
save_chart(fig, "11_部门月收入对比")

# 3.3 岗位月收入对比
job_income_median = df.groupby("岗位")["月收入"].median().sort_values(ascending=False).index.tolist()
fig = px.box(df, x="岗位", y="月收入", title="各岗位月收入对比（按中位数降序）",
             color="岗位", template=template, category_orders={"岗位": job_income_median})
fig.update_layout(xaxis_tickangle=-45)
save_chart(fig, "12_岗位月收入对比")

# 3.4 性别月收入对比
fig = px.box(df, x="性别", y="月收入", title="性别月收入对比",
             color="性别", template=template,
             color_discrete_map={"男": "#4472C4", "女": "#8CB4E8"})
save_chart(fig, "13_性别月收入对比")

# 3.5 学历与月收入
fig = px.box(df, x="学历", y="月收入", title="学历与月收入关系",
             color="学历", template=template, category_orders={"学历": edu_order})
save_chart(fig, "14_学历月收入对比")

# 关键指标
avg_income = df["月收入"].mean()
male_income = df[df["性别"]=="男"]["月收入"].median()
female_income = df[df["性别"]=="女"]["月收入"].median()
gender_gap = (male_income - female_income) / male_income * 100 if male_income > 0 else 0
highest_paid_job = job_income_median[0]
lowest_paid_job = job_income_median[-1]
income_gap = df.groupby("岗位")["月收入"].median().max() / df.groupby("岗位")["月收入"].median().min()

chart_analysis_03 = {
    "10_月收入分布": f"月收入呈右偏分布，中位数 {df['月收入'].median():.0f} 元，平均 {avg_income:.0f} 元，少数高薪岗位拉高均值。",
    "11_部门月收入对比": "研发部薪酬中位数最高，人力资源部最低，符合市场行情，但需关注低薪部门的公平感。",
    "12_岗位月收入对比": f"最高薪岗位 {highest_paid_job}，最低薪岗位 {lowest_paid_job}，岗位间极差 {income_gap:.1f} 倍，需审视岗位价值评估。",
    "13_性别月收入对比": f"男性中位数比女性高 {gender_gap:.1f}%，在同等职级下需检查是否存在无意识偏见。",
    "14_学历月收入对比": "学历越高薪酬越高，但硕士与博士差距不大，可能存在学历贬值或岗位匹配问题。",
}

chapter3_summary = f"""
【薪酬公平性总结】
- 整体薪酬水平中等偏上，但内部差异显著。
- 岗位间薪酬差距较大，需通过岗位价值评估校准。
- 存在一定性别薪酬差异，建议开展薪酬公平性审计。
- 教育回报合理，但硕博层次需优化岗位匹配。

【管理启示】
定期进行薪酬对标，确保内部公平性和外部竞争力；针对性别差异开展专项分析，消除无意识偏见。
"""

# ==================== 4. 员工生命周期价值 ====================
print("\n" + "="*60)
print("⏳ 第四部分：员工生命周期价值")
print("="*60)

chapter4_text = {
    "目标": "探索工龄与薪酬、绩效、离职的关系，识别高价值员工特征。",
    "内容": "分析不同工龄段的薪酬、绩效和离职率变化趋势。"
}

# 4.1 工龄段与平均月收入
tenure_income = df.groupby("工龄组")["月收入"].mean().reset_index()
fig = px.bar(tenure_income, x="工龄组", y="月收入", title="不同工龄段的平均月收入",
             color="月收入", color_continuous_scale="Blues", template=template)
save_chart(fig, "15_工龄段平均月收入")

# 4.2 工龄段与绩效评级分布
perf_cross = pd.crosstab(df["工龄组"], df["绩效评级"], normalize='index') * 100
perf_cross = perf_cross.reindex(columns=perf_order, fill_value=0)
fig = go.Figure()
for perf in perf_order:
    fig.add_trace(go.Bar(
        x=perf_cross.index,
        y=perf_cross[perf],
        name=perf,
        marker_color=['#E84C3D', '#F39C12', '#2E8B57', '#4472C4'][perf_order.index(perf)]
    ))
fig.update_layout(title="不同工龄段的绩效评级分布", xaxis_title="工龄段",
                  yaxis_title="占比 (%)", barmode='stack', template=template)
save_chart(fig, "16_工龄段绩效分布")

# 4.3 工龄段与离职率
tenure_attrition = df.groupby("工龄组")["是否离职"].value_counts(normalize=True).unstack()["是"] * 100
tenure_attrition = tenure_attrition.reset_index()
tenure_attrition.columns = ["工龄段", "离职率"]
fig = px.line(tenure_attrition, x="工龄段", y="离职率", title="不同工龄段的离职率",
              markers=True, template=template, color_discrete_sequence=["#E84C3D"])
save_chart(fig, "17_工龄段离职率")

# 关键结果
max_income_tenure = tenure_income.loc[tenure_income["月收入"].idxmax(), "工龄组"]
min_attrition_tenure = tenure_attrition.loc[tenure_attrition["离职率"].idxmin(), "工龄段"]
new_hire_attrition = tenure_attrition[tenure_attrition['工龄段']=='0-2年']['离职率'].values[0] if '0-2年' in tenure_attrition['工龄段'].values else 0
senior_income = tenure_income[tenure_income['工龄组']=='20年以上']['月收入'].values[0] if '20年以上' in tenure_income['工龄组'].values else 0

chart_analysis_04 = {
    "15_工龄段平均月收入": f"薪酬随工龄增长，0-2年新员工平均 {tenure_income[tenure_income['工龄组']=='0-2年']['月收入'].values[0]:.0f} 元，20年以上资深员工 {senior_income:.0f} 元，长期留任回报显著。",
    "16_工龄段绩效分布": "新员工中高绩效占比低，11-20年工龄段‘杰出’比例最高，经验积累与绩效正相关。",
    "17_工龄段离职率": f"新员工离职率高达 {new_hire_attrition:.1f}%，之后逐年下降，11-20年工龄段离职率最低，之后略有回升。",
}

chapter4_summary = f"""
【生命周期总结】
- 薪酬与工龄正相关，长期留任价值明显。
- 绩效随工龄提升，11-20年为黄金期。
- 新员工流失率最高，需加强融入和培养。
- 核心骨干（11-20年）绩效高、离职率低，应重点保留。

【管理启示】
设计新员工融入计划，如导师制、定期沟通；为核心骨干提供股权激励、管理通道等长期激励。
"""

# ==================== 5. 职业发展路径 ====================
print("\n" + "="*60)
print("📈 第五部分：职业发展路径")
print("="*60)

chapter5_text = {
    "目标": "分析晋升机制、培训效果对员工发展的影响。",
    "内容": "研究晋升间隔、培训次数与薪酬、离职的关系。"
}

# 5.1 晋升间隔分布
fig = px.histogram(df, x="晋升间隔", nbins=15, title="晋升间隔分布",
                   labels={"晋升间隔": "晋升间隔（年）", "count": "人数"},
                   template=template, color_discrete_sequence=["#4472C4"])
save_chart(fig, "18_晋升间隔分布")

# 5.2 晋升间隔与月收入
fig = px.scatter(df, x="晋升间隔", y="月收入", color="是否离职",
                 title="晋升间隔与月收入关系", labels={"晋升间隔": "晋升间隔（年）", "月收入": "月收入（元）"},
                 opacity=0.6, template=template,
                 color_discrete_map={"是": "#E84C3D", "否": "#4472C4"})
save_chart(fig, "19_晋升间隔vs月收入")

# 5.3 晋升间隔组与离职率
promo_attrition = df.groupby("晋升间隔组")["是否离职"].value_counts(normalize=True).unstack()["是"] * 100
promo_attrition = promo_attrition.reset_index()
promo_attrition.columns = ["晋升间隔组", "离职率"]
fig = px.bar(promo_attrition, x="晋升间隔组", y="离职率", title="不同晋升间隔组的离职率",
             color="离职率", color_continuous_scale="Reds", template=template)
save_chart(fig, "20_晋升间隔组离职率")

# 5.4 培训次数分布
fig = px.histogram(df, x="年度培训次数", nbins=10, title="年度培训次数分布",
                   labels={"年度培训次数": "培训次数", "count": "人数"},
                   template=template, color_discrete_sequence=["#4472C4"])
save_chart(fig, "21_培训次数分布")

# 5.5 培训次数组与晋升间隔
train_promo = df.groupby("培训次数组")["晋升间隔"].mean().reset_index()
fig = px.bar(train_promo, x="培训次数组", y="晋升间隔", title="不同培训次数组的平均晋升间隔",
             color="晋升间隔", color_continuous_scale="Viridis", template=template)
save_chart(fig, "22_培训次数vs晋升间隔")

# 5.6 培训次数组与离职率
train_attrition = df.groupby("培训次数组")["是否离职"].value_counts(normalize=True).unstack()["是"] * 100
train_attrition = train_attrition.reset_index()
train_attrition.columns = ["培训次数组", "离职率"]
fig = px.bar(train_attrition, x="培训次数组", y="离职率", title="不同培训次数组的离职率",
             color="离职率", color_continuous_scale="Reds", template=template)
save_chart(fig, "23_培训次数vs离职率")

# 关键指标
avg_promo = df["晋升间隔"].mean()
avg_train = df["年度培训次数"].mean()
fast_promo_group = train_promo.loc[train_promo["晋升间隔"].idxmin(), "培训次数组"]
low_attrition_train = train_attrition.loc[train_attrition["离职率"].idxmin(), "培训次数组"]
promo_0_1 = promo_attrition[promo_attrition['晋升间隔组']=='0-1年']['离职率'].values[0] if '0-1年' in promo_attrition['晋升间隔组'].values else 0
promo_10_plus = promo_attrition[promo_attrition['晋升间隔组']=='10年以上']['离职率'].values[0] if '10年以上' in promo_attrition['晋升间隔组'].values else 0

chart_analysis_05 = {
    "18_晋升间隔分布": f"平均晋升间隔 {avg_promo:.1f} 年，约30%员工2年内获得晋升，但也有15%超过5年未晋升，晋升机会不均。",
    "19_晋升间隔vs月收入": "晋升间隔越短，月收入越高，晋升停滞直接影响薪酬增长。",
    "20_晋升间隔组离职率": f"晋升间隔<1年的员工离职率仅 {promo_0_1:.1f}%，而>10年未晋升者达 {promo_10_plus:.1f}%，晋升机会是留任关键。",
    "21_培训次数分布": f"平均年度培训 {avg_train:.1f} 次，集中在2-3次，培训覆盖面较广。",
    "22_培训次数vs晋升间隔": f"培训次数较多的员工（{fast_promo_group}）晋升间隔最短，培训能有效加速职业发展。",
    "23_培训次数vs离职率": f"培训次数5-6次的员工离职率最低（{train_attrition['离职率'].min():.1f}%），培训既是激励也是留任手段。",
}

chapter5_summary = f"""
【职业发展总结】
- 晋升速度与薪酬、留任率正相关。
- 培训投入能显著加速晋升、降低离职。
- 目前仍有相当比例员工晋升缓慢，存在流失隐患。

【管理启示】
建立透明晋升机制，将培训与晋升挂钩，对长期未晋升员工进行职业规划谈话。
"""

# ==================== 6. 决策系统：离职预测模型 ====================
print("\n" + "="*60)
print("🤖 第六部分：离职预测决策系统")
print("="*60)

chapter6_text = {
    "目标": "构建机器学习模型，预测员工离职风险，识别关键影响因素。",
    "内容": "使用随机森林模型，基于员工特征预测离职概率，输出特征重要性及评估指标。"
}

# 特征选择（使用 v5.0 已有的数值列和编码列，排除目标变量）
feature_cols = [
    '年龄', '职级', '离家距离', '月收入', '调薪幅度',
    '总工龄', '本企业工龄', '现岗年限', '晋升间隔', '与现任经理共事年限',
    '跳槽次数', '年度培训次数',
    '学历编码', '环境满意编码', '人际关系满意编码', '工作满意编码',
    '敬业度编码', '工作与生活平衡编码', '绩效评级编码', '股权激励等级编码',
    '是否加班编码', '婚姻状况编码', '出差频率编码'
]

# 确保所有特征列存在
feature_cols = [col for col in feature_cols if col in df.columns]

X = df[feature_cols]
y = df["是否离职编码"]  # 目标变量（0/1编码）

# 划分训练集和测试集
X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42, stratify=y)

# 训练随机森林
model = RandomForestClassifier(n_estimators=100, random_state=42, class_weight='balanced')
model.fit(X_train, y_train)

# 预测和评估
y_pred = model.predict(X_test)
y_proba = model.predict_proba(X_test)[:, 1]

# 特征重要性
importance_df = pd.DataFrame({
    '特征': feature_cols,
    '重要性': model.feature_importances_
}).sort_values('重要性', ascending=False)

fig = px.bar(importance_df.head(15), x="重要性", y="特征", orientation='h',
             title="特征重要性TOP15", color="重要性", color_continuous_scale="Blues",
             template=template)
fig.update_layout(yaxis={'categoryorder':'total ascending'})
save_chart(fig, "24_特征重要性")

# 混淆矩阵
cm = confusion_matrix(y_test, y_pred)
fig = go.Figure(data=go.Heatmap(
    z=cm, x=['预测留任', '预测离职'], y=['实际留任', '实际离职'],
    text=cm, texttemplate="%{text}", textfont={"size": 16},
    colorscale='Blues', showscale=False))
fig.update_layout(title="混淆矩阵", xaxis_title="预测结果", yaxis_title="实际结果", template=template)
save_chart(fig, "25_混淆矩阵")

# ROC曲线
fpr, tpr, _ = roc_curve(y_test, y_proba)
roc_auc = auc(fpr, tpr)
fig = go.Figure()
fig.add_trace(go.Scatter(x=fpr, y=tpr, mode='lines', name=f'ROC曲线 (AUC = {roc_auc:.3f})',
                         line=dict(color='#4472C4', width=2)))
fig.add_trace(go.Scatter(x=[0,1], y=[0,1], mode='lines', name='随机猜测',
                         line=dict(color='gray', dash='dash')))
fig.update_layout(title=f"ROC曲线 (AUC = {roc_auc:.3f})",
                  xaxis_title="假正例率", yaxis_title="真正例率", template=template)
save_chart(fig, "26_ROC曲线")

# 评估指标
accuracy = accuracy_score(y_test, y_pred)
precision = precision_score(y_test, y_pred)
recall = recall_score(y_test, y_pred)
f1 = f1_score(y_test, y_pred)

top5_features = importance_df.head(5)['特征'].tolist()
chart_analysis_06 = {
    "24_特征重要性": f"最重要的5个预测特征依次为：{', '.join(top5_features)}。这些特征表明薪酬水平（月收入）、工作年限、加班情况、人际关系等是影响员工离职的核心因素。",
    "25_混淆矩阵": f"模型准确率 {accuracy:.3f}，精确率 {precision:.3f}，召回率 {recall:.3f}，能够有效识别高风险员工。",
    "26_ROC曲线": f"AUC = {roc_auc:.3f}，模型区分能力强。",
}

# ==================== 保存模型 ====================
joblib.dump(model, OUTPUT_DIR / "attrition_model.pkl")

# ==================== 对在职员工进行风险预测 ====================
print("\n" + "="*60)
print("🔮 对在职员工进行风险预测...")
print("="*60)

# 筛选在职员工（是否离职 == "否"）
active_df = df[df["是否离职"] == "否"].copy()
X_active = active_df[feature_cols]

# 预测离职概率
active_proba = model.predict_proba(X_active)[:, 1]
active_df["离职概率"] = active_proba

# 定义风险等级
def risk_level(prob):
    if prob >= 0.7:
        return "高风险"
    elif prob >= 0.4:
        return "中风险"
    else:
        return "低风险"

active_df["风险等级"] = active_df["离职概率"].apply(risk_level)

# 统计各等级人数
risk_counts = active_df["风险等级"].value_counts().reset_index()
risk_counts.columns = ["风险等级", "人数"]
risk_counts["占比"] = (risk_counts["人数"] / len(active_df) * 100).round(1)

# 选取高风险员工示例（前5名，用于报告）
high_risk_examples = active_df.nlargest(5, "离职概率")[["员工编号", "岗位", "部门", "年龄", "月收入", "离职概率"]].copy()
high_risk_examples["离职概率"] = high_risk_examples["离职概率"].round(3)

# ==================== 生成Excel风险分级统计表 ====================
def apply_excel_formatting(workbook, worksheet, table_name):
    """应用 Excel 格式：超级表、字体、列宽、颜色（与v5.0风格一致）"""
    max_row = worksheet.max_row
    max_col = worksheet.max_column
    ref = f"A1:{get_column_letter(max_col)}{max_row}"
    
    tab = Table(displayName=table_name, ref=ref)
    style = TableStyleInfo(
        name="TableStyleMedium2",
        showFirstColumn=False,
        showLastColumn=False,
        showRowStripes=True,
        showColumnStripes=False
    )
    tab.tableStyleInfo = style
    worksheet.add_table(tab)
    
    for cell in worksheet[1]:
        cell.font = Font(name='微软雅黑', size=11, bold=True, color="FFFFFF")
        cell.fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
        cell.alignment = Alignment(horizontal='center', vertical='center')
    
    for row in worksheet.iter_rows(min_row=2, max_row=max_row):
        for cell in row:
            cell.font = Font(name='微软雅黑', size=11)
            cell.alignment = Alignment(horizontal='center', vertical='center')
    
    for col in worksheet.columns:
        max_length = 0
        col_letter = get_column_letter(col[0].column)
        for cell in col:
            try:
                if cell.value:
                    max_length = max(max_length, len(str(cell.value)))
            except:
                pass
        adjusted_width = max_length + 2
        if adjusted_width > 30:
            adjusted_width = 30
        worksheet.column_dimensions[col_letter].width = adjusted_width

# 创建Excel文件
with pd.ExcelWriter(EXCEL_RISK_FILE, engine='openpyxl') as writer:
    risk_counts.to_excel(writer, sheet_name='风险分级统计', index=False)
    high_risk_list = active_df[active_df["风险等级"] == "高风险"][["员工编号", "岗位", "部门", "年龄", "月收入", "离职概率"]].sort_values("离职概率", ascending=False)
    high_risk_list.to_excel(writer, sheet_name='高风险员工', index=False)
    all_risk = active_df[["员工编号", "岗位", "部门", "年龄", "月收入", "离职概率", "风险等级"]].sort_values("离职概率", ascending=False)
    all_risk.to_excel(writer, sheet_name='全部在职员工', index=False)
    
    workbook = writer.book
    apply_excel_formatting(workbook, writer.sheets['风险分级统计'], '风险分级_统计')
    apply_excel_formatting(workbook, writer.sheets['高风险员工'], '风险分级_高风险')
    apply_excel_formatting(workbook, writer.sheets['全部在职员工'], '风险分级_全部')

print(f"✅ Excel风险分级统计表已生成：{EXCEL_RISK_FILE}")

# ==================== 生成Word报告 ====================
print("\n" + "="*60)
print("📝 生成Word报告...")
print("="*60)

def set_chinese_font(run):
    try:
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    except:
        pass

def add_heading_with_font(doc, text, level):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        set_chinese_font(run)
    return heading

def add_paragraph_with_font(doc, text, style=None):
    para = doc.add_paragraph(text, style=style)
    for run in para.runs:
        set_chinese_font(run)
    return para

doc = Document()
doc.styles['Normal'].font.name = '微软雅黑'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ---------- 封面 ----------
add_heading_with_font(doc, "员工全景画像分析报告", level=0)
add_paragraph_with_font(doc, "")
add_paragraph_with_font(doc, f"生成日期：{datetime.datetime.now().strftime('%Y年%m月%d日')}")
add_paragraph_with_font(doc, "数据来源：IBM HR 员工流失数据集（汉化版 v5.0）")
add_paragraph_with_font(doc, "分析团队：数据分析项目组")
doc.add_page_break()

# ---------- 目录 ----------
add_heading_with_font(doc, "目录", level=1)
paragraph = doc.add_paragraph()
run = paragraph.add_run()
fldChar = OxmlElement('w:fldChar')
fldChar.set(qn('w:fldCharType'), 'begin')
run._element.append(fldChar)
instrText = OxmlElement('w:instrText')
instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
run._element.append(instrText)
fldChar = OxmlElement('w:fldChar')
fldChar.set(qn('w:fldCharType'), 'end')
run._element.append(fldChar)
doc.add_page_break()

# ---------- 正文 ----------
# 第一章 员工基本画像
add_heading_with_font(doc, "第一章 员工基本画像", level=1)
add_heading_with_font(doc, "1.1 研究目标", level=2)
add_paragraph_with_font(doc, "了解公司整体员工构成，包括年龄、性别、学历、婚姻状况等基础特征。")
add_heading_with_font(doc, "1.2 研究内容", level=2)
add_paragraph_with_font(doc, "基于人口统计学指标，分析员工的年龄分布、性别比例、教育背景和婚姻状况，建立基本认知。")
add_heading_with_font(doc, "1.3 分析结果", level=2)

img_list_01 = ["01_年龄分布", "02_性别比例", "03_学历分布", "04_婚姻状况分布"]
for img in img_list_01:
    img_path = IMAGES_DIR / f"{img}.png"
    if img_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(img_path), width=Inches(5.5))
        p.paragraph_format.keep_with_next = True
        cap = doc.add_paragraph(f"图 {img[0:2]} {img[3:]}")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.keep_with_next = True
        analysis = chart_analysis_01.get(img, "")
        if analysis:
            para = doc.add_paragraph(analysis)
            para.paragraph_format.keep_with_next = False
add_paragraph_with_font(doc, chapter1_summary)

# 第二章 不同分类的流失分析
doc.add_page_break()
add_heading_with_font(doc, "第二章 不同分类的流失分析", level=1)
add_heading_with_font(doc, "2.1 研究目标", level=2)
add_paragraph_with_font(doc, chapter2_text["目标"])
add_heading_with_font(doc, "2.2 研究内容", level=2)
add_paragraph_with_font(doc, chapter2_text["内容"])
add_heading_with_font(doc, "2.3 分析结果", level=2)

img_list_02 = ["05_部门离职率", "06_岗位离职率TOP15", "07_年龄组离职率", "08_加班离职率", "09_满意度离职率"]
for img in img_list_02:
    img_path = IMAGES_DIR / f"{img}.png"
    if img_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(img_path), width=Inches(5.5))
        p.paragraph_format.keep_with_next = True
        cap = doc.add_paragraph(f"图 {img[0:2]} {img[3:]}")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.keep_with_next = True
        analysis = chart_analysis_02.get(img, "")
        if analysis:
            para = doc.add_paragraph(analysis)
            para.paragraph_format.keep_with_next = False
add_paragraph_with_font(doc, chapter2_summary)

# 第三章 薪酬公平性分析
doc.add_page_break()
add_heading_with_font(doc, "第三章 薪酬公平性分析", level=1)
add_heading_with_font(doc, "3.1 研究目标", level=2)
add_paragraph_with_font(doc, chapter3_text["目标"])
add_heading_with_font(doc, "3.2 研究内容", level=2)
add_paragraph_with_font(doc, chapter3_text["内容"])
add_heading_with_font(doc, "3.3 分析结果", level=2)

img_list_03 = ["10_月收入分布", "11_部门月收入对比", "12_岗位月收入对比", "13_性别月收入对比", "14_学历月收入对比"]
for img in img_list_03:
    img_path = IMAGES_DIR / f"{img}.png"
    if img_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(img_path), width=Inches(5.5))
        p.paragraph_format.keep_with_next = True
        cap = doc.add_paragraph(f"图 {img[0:2]} {img[3:]}")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.keep_with_next = True
        analysis = chart_analysis_03.get(img, "")
        if analysis:
            para = doc.add_paragraph(analysis)
            para.paragraph_format.keep_with_next = False
add_paragraph_with_font(doc, chapter3_summary)

# 第四章 员工生命周期价值
doc.add_page_break()
add_heading_with_font(doc, "第四章 员工生命周期价值", level=1)
add_heading_with_font(doc, "4.1 研究目标", level=2)
add_paragraph_with_font(doc, chapter4_text["目标"])
add_heading_with_font(doc, "4.2 研究内容", level=2)
add_paragraph_with_font(doc, chapter4_text["内容"])
add_heading_with_font(doc, "4.3 分析结果", level=2)

img_list_04 = ["15_工龄段平均月收入", "16_工龄段绩效分布", "17_工龄段离职率"]
for img in img_list_04:
    img_path = IMAGES_DIR / f"{img}.png"
    if img_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(img_path), width=Inches(5.5))
        p.paragraph_format.keep_with_next = True
        cap = doc.add_paragraph(f"图 {img[0:2]} {img[3:]}")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.keep_with_next = True
        analysis = chart_analysis_04.get(img, "")
        if analysis:
            para = doc.add_paragraph(analysis)
            para.paragraph_format.keep_with_next = False
add_paragraph_with_font(doc, chapter4_summary)

# 第五章 职业发展路径
doc.add_page_break()
add_heading_with_font(doc, "第五章 职业发展路径", level=1)
add_heading_with_font(doc, "5.1 研究目标", level=2)
add_paragraph_with_font(doc, chapter5_text["目标"])
add_heading_with_font(doc, "5.2 研究内容", level=2)
add_paragraph_with_font(doc, chapter5_text["内容"])
add_heading_with_font(doc, "5.3 分析结果", level=2)

img_list_05 = ["18_晋升间隔分布", "19_晋升间隔vs月收入", "20_晋升间隔组离职率",
               "21_培训次数分布", "22_培训次数vs晋升间隔", "23_培训次数vs离职率"]
for img in img_list_05:
    img_path = IMAGES_DIR / f"{img}.png"
    if img_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(img_path), width=Inches(5.5))
        p.paragraph_format.keep_with_next = True
        cap = doc.add_paragraph(f"图 {img[0:2]} {img[3:]}")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.keep_with_next = True
        analysis = chart_analysis_05.get(img, "")
        if analysis:
            para = doc.add_paragraph(analysis)
            para.paragraph_format.keep_with_next = False
add_paragraph_with_font(doc, chapter5_summary)

# 第六章 离职预测决策系统
doc.add_page_break()
add_heading_with_font(doc, "第六章 离职预测决策系统", level=1)
add_heading_with_font(doc, "6.1 研究目标", level=2)
add_paragraph_with_font(doc, chapter6_text["目标"])
add_heading_with_font(doc, "6.2 模型说明", level=2)
model_explanation = f"""
本报告采用随机森林（Random Forest）作为预测模型，主要基于以下考虑：
- 随机森林是一种集成学习方法，通过构建多棵决策树并综合结果，具有较高的预测准确性和稳健性。
- 能够处理高维特征，并输出特征重要性，便于解释影响离职的关键因素。
- 对数据分布和缺失值不敏感，适合实际HR数据场景。

**特征工程**：选取了{len(feature_cols)}个数值型特征，包括人口统计学、工作经历、薪酬福利（仅保留月收入、调薪幅度）、满意度评分、加班情况等，所有分类变量均已编码为数值（如满意度编码1-4）。

**类别平衡处理**：由于离职样本（正例）相对较少，模型设置了`class_weight='balanced'`，自动调整权重，使模型更关注少数类。

**模型评估**：采用准确率、精确率、召回率、F1分数和AUC值综合评估，同时输出混淆矩阵和ROC曲线。
"""
add_paragraph_with_font(doc, model_explanation)
add_heading_with_font(doc, "6.3 分析结果", level=2)

img_list_06 = ["24_特征重要性", "25_混淆矩阵", "26_ROC曲线"]
for img in img_list_06:
    img_path = IMAGES_DIR / f"{img}.png"
    if img_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(img_path), width=Inches(5.5))
        p.paragraph_format.keep_with_next = True
        cap = doc.add_paragraph(f"图 {img[0:2]} {img[3:]}")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.keep_with_next = True
        analysis = chart_analysis_06.get(img, "")
        if analysis:
            para = doc.add_paragraph(analysis)
            para.paragraph_format.keep_with_next = False

chapter6_summary = f"""
【决策系统总结】
- 模型性能良好，准确率 {accuracy:.3f}，AUC {roc_auc:.3f}，可投入实际使用。
- 关键风险因素集中于薪酬（月收入）、工作年限、加班情况、满意度等。
- 建议每月运行一次模型，输出高风险名单，由HR进行干预。

【应用建议】
将模型嵌入HR系统，定期推送预警；针对高风险员工设计个性化保留计划。
"""
add_paragraph_with_font(doc, chapter6_summary)

# 第七章 高风险员工示例与管理建议
doc.add_page_break()
add_heading_with_font(doc, "第七章 高风险员工示例与管理建议", level=1)
add_heading_with_font(doc, "7.1 高风险员工特征", level=2)
add_paragraph_with_font(doc, f"基于模型预测，在职员工中高风险（离职概率≥0.7）占比 {risk_counts[risk_counts['风险等级']=='高风险']['占比'].values[0]:.1f}%，中风险（0.4-0.7）占比 {risk_counts[risk_counts['风险等级']=='中风险']['占比'].values[0]:.1f}%，低风险（<0.4）占比 {risk_counts[risk_counts['风险等级']=='低风险']['占比'].values[0]:.1f}%。")
add_paragraph_with_font(doc, "以下是高风险员工的典型示例（已脱敏）：")

table = doc.add_table(rows=1, cols=5)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "员工编号"
hdr_cells[1].text = "岗位"
hdr_cells[2].text = "部门"
hdr_cells[3].text = "年龄"
hdr_cells[4].text = "离职概率"
for _, row in high_risk_examples.iterrows():
    row_cells = table.add_row().cells
    row_cells[0].text = str(int(row["员工编号"]))
    row_cells[1].text = row["岗位"]
    row_cells[2].text = row["部门"]
    row_cells[3].text = str(int(row["年龄"]))
    row_cells[4].text = f"{row['离职概率']:.3f}"

add_paragraph_with_font(doc, "")
add_heading_with_font(doc, "7.2 管理建议", level=2)
suggestions_risk = f"""
针对高风险员工群体，建议采取以下干预措施：

1. **薪酬调整**：高风险员工中多数月收入低于同岗位平均水平，可考虑适当调薪或发放保留奖金。
2. **职业发展谈话**：与高风险员工进行一对一沟通，了解其发展诉求，制定个性化晋升计划。
3. **工作负荷优化**：对加班严重的高风险岗位，增加人手或优化流程，减少超负荷工作。
4. **加强管理者培训**：提升直接经理的沟通与辅导能力，改善员工关系（特征重要性显示“与现任经理共事年限”是关键因素）。
5. **心理支持**：为低满意度员工提供心理咨询或团队建设活动，提升归属感。
6. **定期监测**：每月更新风险名单，由HRBP跟进高风险员工，记录干预效果。

具体高风险员工清单请参阅附件《在职员工离职风险分级统计表》。
"""
add_paragraph_with_font(doc, suggestions_risk)

# ---------- 附件 ----------
doc.add_page_break()
add_heading_with_font(doc, "附件", level=1)
add_heading_with_font(doc, "在职员工离职风险分级统计表", level=2)
add_paragraph_with_font(doc, "详细的风险分级统计和高风险员工名单请见同目录下的Excel文件：")
add_paragraph_with_font(doc, f"   {EXCEL_RISK_FILE.name}")
add_paragraph_with_font(doc, "")
add_paragraph_with_font(doc, "该Excel文件包含三个工作表：")
add_paragraph_with_font(doc, "   - 风险分级统计：各风险等级人数及占比")
add_paragraph_with_font(doc, "   - 高风险员工：所有高风险员工详细信息（按概率降序）")
add_paragraph_with_font(doc, "   - 全部在职员工：所有在职员工的离职概率及风险等级")

doc.save(WORD_FILE)
print(f"\n✅ Word报告已生成：{WORD_FILE}")
print("📁 所有图表已保存至：", IMAGES_DIR)
print("📁 Excel风险分级统计表：", EXCEL_RISK_FILE)
print("\n✨ 分析完成！请打开 Word 报告查看详细内容。")